import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_inline(paragraph, text):
    """
    Parses inline markdown syntax for bold (**) and italic (*).
    Supports mixed formats.
    """
    # Pattern to find **bold** or *italic* patterns
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
        else:
            paragraph.add_run(part)

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Configure document styles for a professional look
    # Normal Style (body text)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)  # Soft charcoal
    
    # Adjust spacing
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line_str = line.strip()
        
        # Handle empty lines
        if not line_str:
            continue
            
        # Horizontal rules
        if line_str == '---' or line_str == '***':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("—" * 50)
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            continue
            
        # Headings (up to level 4)
        h_match = re.match(r'^(#{1,6})\s+(.*)$', line_str)
        if h_match:
            level = len(h_match.group(1))
            title_text = h_match.group(2)
            
            # Since Heading 1 and 2 should look distinct and prominent:
            p = doc.add_heading(level=level)
            p.paragraph_format.keep_with_next = True
            
            # Custom font color and sizes for headings
            run = p.add_run(title_text)
            run.font.name = 'Calibri Light'
            
            if level == 1:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Navy Blue
                run.bold = True
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
            elif level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)  # Mid-Tone Steel Blue
                run.bold = True
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
            elif level == 3:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x3B, 0x6E, 0x6E)  # Muted Teal
                run.bold = True
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            else:
                run.font.size = Pt(11.5)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)  # Dark Grey
                run.bold = True
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
            continue
            
        # Bullet list items
        bullet_match = re.match(r'^[\*\-\+]\s+(.*)$', line_str)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_inline(p, bullet_match.group(1))
            continue
            
        # Numbered list items
        num_match = re.match(r'^(\d+)\.\s+(.*)$', line_str)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            parse_inline(p, num_match.group(2))
            continue
            
        # Standard paragraph
        p = doc.add_paragraph()
        parse_inline(p, line_str)
        
    doc.save(docx_path)
    print(f"Successfully generated docx: {docx_path}")

if __name__ == "__main__":
    md_file = "presentation_key_concepts.md"
    docx_file = "presentation_key_concepts.docx"
    md_to_docx(md_file, docx_file)
